from __future__ import annotations
import os, hashlib, datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordWriter:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        if os.path.exists(docx_path):
            self.doc = Document(docx_path)
        else:
            self.doc = Document()
            self._init_template()
            self.save()

    def _init_template(self):
        self.doc.add_heading('Test Evidence Report', level=0)
        self.doc.add_paragraph(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        self.doc.add_heading('Test Steps', level=1); self.doc.add_paragraph('')
        self.doc.add_heading('Screenshots', level=1); self.doc.add_paragraph('')
        self.doc.add_heading('Notes/Observations', level=1); self.doc.add_paragraph('')
        self.doc.add_heading('Issues', level=1); self.doc.add_paragraph('')

    def save(self):
        self.doc.save(self.docx_path)

    def append_steps(self, steps: list[str]):
        self.doc.add_heading('Test Steps', level=1)
        for s in steps:
            p = self.doc.add_paragraph()
            run = p.add_run("Step: ")
            run.bold = True
            p.add_run(s)
            p.add_run(f"  [{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]")
        self.save()

    def add_image_with_caption(self, image_path: str, caption: str, max_width_px: int = 1200):
        self.doc.add_heading('Screenshots', level=1)
        from PIL import Image
        img = Image.open(image_path)
        max_width_inches = max_width_px / 96.0
        try:
            self.doc.add_picture(image_path, width=Inches(max_width_inches))
        except Exception:
            self.doc.add_picture(image_path)
        sha256 = hashlib.sha256(open(image_path, 'rb').read()).hexdigest()
        cap = self.doc.add_paragraph(f"Fig: {caption} — {os.path.basename(image_path)}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hash_p = self.doc.add_paragraph(f"SHA-256: {sha256}")
        hash_p.runs[0].font.size = Pt(8)
        self.save()
        return sha256
